######################################################################
# “report_generator”
######################################################################
from docx import Document
from docx.shared import Inches
from io import BytesIO
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import pandas as pd

def add_heading(document, text, level=1):
    document.add_heading(text, level=level)

def add_paragraph(document, text):
    document.add_paragraph(text)

def set_table_all_borders(table):
    """
    Add borders on all sides and between cells for a python-docx table.
    """
    tbl = table._element
    # ensure tblPr exists
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # remove existing borders if present
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    # create new borders element
    borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{border_name}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')        # 4 = 1/8 point
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '000000')
        borders.append(elem)

    tblPr.append(borders)

def add_dataframe(document, df, title=None):
    if title:
        add_heading(document, title, level=2)

    # create table
    table = document.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)

    # fill rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # apply full borders
    set_table_all_borders(table)

def add_plot(document, fig, title):
    add_heading(document, title, level=2)
    image_stream = BytesIO()
    fig.savefig(image_stream, format='png', bbox_inches='tight')
    image_stream.seek(0)
    # center the image
    p = document.add_paragraph()
    p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_stream, width=Inches(5.5))
    plt.close(fig)

def generate_report(pipeline_results):
    document = Document()

    # Title
    document.add_heading('Probabilistic Decline Curve Analysis Report', 0)

    # Section 1: Overview
    add_heading(document, "1. Overview", level=1)
    add_paragraph(document, (
        "This report summarizes the results of the probabilistic decline curve analysis, "
        "including data cleaning, outlier detection, Monte Carlo sampling, model fitting, "
        "hindcast testing, and estimated ultimate recovery (EUR) analysis."
    ))
    add_paragraph(document, (
        "Four models were fit to each synthetic sample:\n"
        "- **Arps** (Exponential/Hyperbolic)\n"
        "- **Stretched Exponential Model (SEM)**\n"
        "- **Logistic Growth Model (LGM)**\n"
        "- **Capacitance-Resistance Model (CRM)**"
    ))

    # Section 2: Initial Data & Outlier Detection
    add_heading(document, "2. Initial Production Data & Outlier Detection", level=1)
    add_plot(document, pipeline_results['lof_plot'], "Initial Production Data with LOF Outlier Detection")

    # Section 2a: Monte Carlo Sampling
    add_heading(document, "2a. Monte Carlo Sampling", level=2)
    add_plot(document, pipeline_results['sample_fig'], "Sampled N Sorted Data Sets")

    # Section 3: Model Probabilities
    add_heading(document, "3. Marginal Posterior Probabilities of Models", level=1)
    add_plot(document, pipeline_results['prob_plot'], "Model Posterior Probabilities")

    # Section 3a: Training Fit Plots
    add_heading(document, "3a. Training Fit per Model", level=2)
    for model_name, fig in pipeline_results['train_fits'].items():
        add_plot(document, fig, f"Training Fit — {model_name.upper()}")

    # Section 3b: Hindcast Test Plots
    add_heading(document, "3b. Hindcast Test per Model", level=2)
    for model_name, fig in pipeline_results['hindcast_plots'].items():
        add_plot(document, fig, f"Hindcast Test — {model_name.upper()}")

    # Section 4: Model-specific EUR
    add_heading(document, "4. Model-Specific EUR Statistics", level=1)
    model_eur_df = pd.DataFrame(pipeline_results['model_eur_stats']).T
    model_eur_df = model_eur_df.applymap(lambda x: f"{int(round(x, 0)):,}")
    desired_columns = ["p10", "p50", "mean", "p90"]
    model_eur_df_clean = model_eur_df[desired_columns].reset_index().rename(columns={'index': 'Model'})
    add_dataframe(document, model_eur_df_clean, title="Per Model EUR Summary")

    # Section 5: Combined EUR
    add_heading(document, "5. Combined EUR Statistics", level=1)
    combined_eur_stats_df = pd.DataFrame([pipeline_results['combined_eur_stats']])
    combined_eur_stats_df_clean = combined_eur_stats_df[desired_columns].applymap(lambda x: f"{int(round(x, 0)):,}")
    add_dataframe(document, combined_eur_stats_df_clean, title="Combined Model EUR Summary")

    # Section 5a: EUR Boxplot
    add_heading(document, "5a. Multi‑Model EUR Boxplot", level=2)
    add_plot(document, pipeline_results['eur_plot'], "Boxplot of Multimodel Probabilistic EUR")

    # Conclusion
    add_heading(document, "6. Conclusion", level=1)
    add_paragraph(document, (
        "The analysis demonstrates the range of production forecasts and uncertainties "
        "associated with the selected decline curve models. Multi‑model probabilistic "
        "forecasts provide a robust outlook for future production."
    ))

    # Save document to BytesIO
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io
